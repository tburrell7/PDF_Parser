from detail import Detail, DetailType
from docx import Document

def export(details: list[Detail], file_name: str):
    procedure_details = list[Detail]
    medication_details = list[Detail]
    allergy_details = list[Detail]
    for detail in details:
    # not using match because using 3.9
        if detail.type == DetailType.PROCEDURE:
            procedure_details.append(detail)
        elif detail.type == DetailType.MEDICATION:
            medication_details.append(detail)
        elif detail.type == DetailType.ALLERGY:
            allergy_details.append(detail)
        else:
            print(f"unknown detail type {detail.type}")


    doc = Document()

    doc.add_heading("Procedures", 2)
    for detail in procedure_details:
        doc.add_paragraph(f"Name: {detail.value}\t Page: {detail.page}", "List Bullet")

    doc.add_heading("Medications", 2)
    # TODO: add dates
    for detail in medication_details:
        doc.add_paragraph(f"Name: {detail.value}\t Page: {detail.page}", "List Bullet")

    doc.add_heading("Allergies", 2)
    for detail in allergy_details:
        doc.add_paragraph(f"Name: {detail.value}\t Page: {detail.page}", "List Bullet")

    doc.save(file_name)
    # TODO check for existence
            