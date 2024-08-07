from detail import Detail, DetailType
from docx import Document

def export(details: list[Detail], file_name: str = "output.docx"):
    print("Creating output")
    procedure_details: list[Detail] = []
    medication_details: list[Detail] = []
    allergy_details: list[Detail] = []
    for detail in details:
        if detail.type == DetailType.PROCEDURE:
            procedure_details.append(detail)
        elif detail.type == DetailType.MEDICATION:
            medication_details.append(detail)
        elif detail.type == DetailType.ALLERGY:
            allergy_details.append(detail)
        else:
            print(f"unknown detail type {detail.type}")


    doc = Document()
    doc.add_heading(f"{file_name} Patient Details", 1)

    doc.add_heading("Procedures", 2)
    for detail in procedure_details:
        if detail.date is not None:
            doc.add_paragraph(f"\tProcedure {detail.value} occured on Date {detail.date}. Referenced on Page: {detail.page}", "List Number")
        else:
            doc.add_paragraph(f"\tProcedure {detail.value}. Referenced on Page: {detail.page}", "List Number")

    doc.add_heading("Medications", 2)
    for detail in medication_details:
        if detail.date is not None:
            doc.add_paragraph(f"\tMedication {detail.value} occured on Date {detail.date}. Referenced on Page: {detail.page}", "List Number")
        else:
            doc.add_paragraph(f"\tMedication {detail.value}. Referenced on Page: {detail.page}", "List Number")

    doc.add_heading("Allergies", 2)
    for detail in allergy_details:
        doc.add_paragraph(f"\t Allergy {detail.value}. Referenced on Page: {detail.page}", "List Number")

    doc.save("output.docx")
            